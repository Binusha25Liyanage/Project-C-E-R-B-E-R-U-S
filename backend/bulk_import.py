"""
Bulk import: read many spreadsheet/Word-table files that share one structure,
map their columns to a schema's fields once, and import all rows through the
same validate_record() pipeline used by manual entry - so a bulk-imported
row and a hand-typed row are held to identical rules.
"""

import os
import csv
import openpyxl
from docx import Document as DocxDocument

from .validation import validate_record


def read_file_columns(file_path):
    """Return the header row of a file, for the column-mapping UI."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".csv":
        with open(file_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            return next(reader, [])
    elif ext in (".xlsx", ".xlsm"):
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        ws = wb.active
        first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), ())
        return [str(c) if c is not None else "" for c in first_row]
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        if not doc.tables:
            return []
        header_row = doc.tables[0].rows[0]
        return [cell.text.strip() for cell in header_row.cells]
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _read_file_rows(file_path):
    """Yield each data row (dict keyed by the file's own header names)."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".csv":
        with open(file_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                yield row

    elif ext in (".xlsx", ".xlsm"):
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        headers = [str(c) if c is not None else "" for c in next(rows_iter, ())]
        for raw_row in rows_iter:
            if all(v is None for v in raw_row):
                continue  # skip fully blank rows
            yield {headers[i]: raw_row[i] for i in range(len(headers)) if i < len(raw_row)}

    elif ext == ".docx":
        doc = DocxDocument(file_path)
        if not doc.tables:
            return
        table = doc.tables[0]
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            yield {headers[i]: values[i] for i in range(len(headers)) if i < len(values)}

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def run_bulk_import(db, job_id, schema, file_paths, column_mapping, batch_label, progress_cb):
    """
    column_mapping: { schema_field_name: source_column_name_or_None }
    progress_cb(processed_files, total_files, current_file, rows_done, rows_total)
    Returns a result dict summarizing the whole batch.
    """
    fields = schema["fields"]
    file_results = []
    total_imported = 0

    for file_index, file_path in enumerate(file_paths):
        filename = os.path.basename(file_path)
        file_errors = []
        imported = 0
        row_num = 1  # header is row 1

        try:
            rows = list(_read_file_rows(file_path))
        except Exception as e:
            file_results.append({
                "file": filename, "total": 0, "imported": 0, "failed": 0,
                "errors": [{"row": None, "message": f"Could not read file: {e}"}],
            })
            progress_cb(file_index + 1, len(file_paths), filename, 0, 0)
            continue

        total_rows = len(rows)
        for i, source_row in enumerate(rows):
            row_num = i + 2  # +1 for header, +1 for 1-indexing
            mapped_data = {}
            for field in fields:
                source_col = column_mapping.get(field["name"])
                mapped_data[field["name"]] = source_row.get(source_col) if source_col else None

            is_valid, errors, cleaned = validate_record(fields, mapped_data)
            if is_valid:
                db.create_record(schema["id"], cleaned, source="bulk_import", status="pending")
                imported += 1
                total_imported += 1
            else:
                file_errors.append({"row": row_num, "message": "; ".join(f"{k}: {v}" for k, v in errors.items())})

            if i % 10 == 0 or i == total_rows - 1:
                progress_cb(file_index, len(file_paths), filename, i + 1, total_rows)

        file_results.append({
            "file": filename,
            "total": total_rows,
            "imported": imported,
            "failed": len(file_errors),
            "errors": file_errors[:50],  # cap so a badly-mapped file doesn't flood the UI
        })
        progress_cb(file_index + 1, len(file_paths), filename, total_rows, total_rows)

    return {
        "batch_label": batch_label,
        "files": file_results,
        "total_imported": total_imported,
        "total_files": len(file_paths),
    }
